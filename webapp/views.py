import io
import subprocess
import urllib.request
from tempfile import NamedTemporaryFile

from groq import Groq
import yt_dlp
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt
from ibm_watson import SpeechToTextV1
from ibm_cloud_sdk_core.authenticators import IAMAuthenticator
from django.http import JsonResponse
import os
import PyPDF2
import docx
from django.template.loader import render_to_string
from django.utils.html import escape

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt
import json

import base64
import re

# import textract
# from docx import Document
# WatsonX API setup

WATSON_API_KEY = os.environ.get('WATSON_API_KEY', 'DyzopmVFpgPTJi-YFnXvJsxK1U1Dk6WIZA8a2V3IGYqx')
WATSON_URL = os.environ.get('WATSON_URL',
                            'https://api.au-syd.speech-to-text.watson.cloud.ibm.com/instances/b088efec-be09-4d11-8b27-09a420dc010c')

authenticator = IAMAuthenticator(WATSON_API_KEY)
speech_to_text = SpeechToTextV1(authenticator=authenticator)
speech_to_text.set_service_url(WATSON_URL)


def transcribe(request):
    return render(request, 'Transcribing.html')


def convert_to_wav(input_file):
    """
    Converts any audio/video file to 16kHz WAV format for WatsonX.
    Args:
    input_file: The input file path.
    Returns:
    The path to the converted WAV file.
    """
    output_file = os.path.join(settings.MEDIA_ROOT, 'converted_audio.wav')
    try:
        # Convert the file to 16kHz WAV format suitable for WatsonX
        subprocess.run(
            ['ffmpeg', '-i', input_file, '-ar', '16000', '-ac', '1', output_file],
            check=True
        )
        return output_file
    except subprocess.CalledProcessError as e:
        print(f"Error during conversion: {e}")
        return None


def download_youtube_video_as_wav(url):
    """
    Downloads the YouTube video directly as a WAV file.
    Args:
    url: The YouTube URL.
    Returns:
    The path to the downloaded WAV file.
    """
    output_file = os.path.join(settings.MEDIA_ROOT, 'youtube_audio')
    ydl_opts = {
        'format': 'bestaudio/best',
        'postprocessors': [{
            'key': 'FFmpegExtractAudio',
            'preferredcodec': 'wav',
            'preferredquality': '192',
            'nopostoverwrites': True
        }],
        'outtmpl': output_file
    }
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([url])
        return output_file
    except Exception as e:
        print(f"Error downloading YouTube video: {e}")
        return None


def transcribe_audio(file_path):
    """
    Transcribes an audio file using WatsonX Speech to Text service.
    Args:
    file_path: The path to the audio file.
    Returns:
    The transcription result.
    """
    try:
        print("Transcribing audio...")
        with open(file_path, 'rb') as audio_file:
            print('file path : in transcribe function', file_path)
            result = speech_to_text.recognize(
                audio=audio_file,
                content_type='audio/wav',
                model='en-US_BroadbandModel'
            ).get_result()
            print("resultsss", result)
        return result
    except Exception as e:
        print(f"Error during transcription: {e}")
        return None


transcribed_saved_text = ""


def trans_results(request):
    global transcribed_saved_text
    if request.method == 'POST':
        youtube_url = request.POST.get('url')
        uploaded_file = request.FILES.get('file')
        chat_query = request.POST.get('query')

        output_file_path = None
        fs = FileSystemStorage()

        # Handle file upload
        if uploaded_file:
            # Save uploaded file temporarily
            file_path = fs.save(uploaded_file.name, uploaded_file)
            file_full_path = fs.path(file_path)
            print(f"Uploaded file saved at {file_full_path}")

            # Convert uploaded file to WAV format
            output_file_path = convert_to_wav(file_full_path)

        # Handle YouTube URL input
        elif youtube_url:
            # Download YouTube video as WAV
            print(f"Downloading video from YouTube: {youtube_url}")
            output_file_path = download_youtube_video_as_wav(youtube_url)
            output_file_path += ".wav"
            print(f"WAV file saved at ", os.path.exists(output_file_path))

        # If conversion/download was successful, send the audio to WatsonX for transcription
        if output_file_path and os.path.exists(output_file_path):
            print("Output file path conditions")
            result = transcribe_audio(output_file_path)
            if result is None:
                result = "No transcription because of no internet connection"
                os.remove(output_file_path)
                return render(request, 'Transcribe-2.html', {'transcription': result})

            # Process transcription text
            transcribe = [text["alternatives"][0]['transcript'].rstrip() + '.\n' for text in result['results']]
            transcribed_text = ''.join(transcribe)
            transcribed_saved_text = transcribed_text
            os.remove(output_file_path)

            # Return transcribed text if no query is provided
            if chat_query is None:
                return render(request, 'Transcribe-2.html', {'transcription': transcribed_text})

            # If a query is present, use transcribed text and query to generate response

        elif transcribed_saved_text:
            bot_response = generate_prompt(transcribed_saved_text, chat_query)
            return JsonResponse({'message': bot_response,
                                 'status': 'success'})
        elif chat_query is not None:
            # Default response if no transcription but only query is provided
            bot_response = "No transcription context available. How can I assist you with your query?"
            data = {
                "message": bot_response,
                'status': 'success'
            }
            return JsonResponse(data)

    return render(request, 'Transcribe-2.html')


def generate_prompt(transcribed_text, user_query):
    client = Groq(
        api_key="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",
    )

    # Define the prompt with context and user query
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": f"You are a helpful assistant. Here is the context: '{transcribed_text}'."
            },
            {
                "role": "user",
                "content": user_query
            }
        ],
        model="llama3-8b-8192",
        temperature=0.5,
        max_tokens=1024,
        top_p=1,
        stop=None,
        stream=False
    )

    # Return the assistant's response
    return chat_completion.choices[0].message.content


# Initialize context as a list to store text from uploaded files
context_list = []


def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
    return text


def extract_text_from_docx(file_path):
    text = ""
    doc = docx.Document(file_path)
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    print(text)
    return text


def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()


def generate_prompt2(context,user_query):
    # Assuming your Groq API key and client initialization here...
    client = Groq(
        api_key="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",
    )

    # Define the prompt with context and user query
    chat_completion = client.chat.completions.create(
        messages=[
            {
                "role": "system",
                "content": f"You are a helpful assistant. and here is the context {context}"
            },
            {
                "role": "user",
                "content": user_query
            }
        ],
        model="llama3-8b-8192",
        temperature=0.5,
        max_tokens=1024,
        top_p=1,
        stop=None,
        stream=False
    )

    # Return the assistant's response
    return chat_completion.choices[0].message.content


def chat_with_book(request):
    global context_list  # Use the global context list to store extracted text

    if request.method == 'POST':
        chat_query = request.POST.get('query')

        # Handle file uploads
        uploaded_file = request.FILES.get('file')
        print("File",uploaded_file)
        if uploaded_file:
            fs = FileSystemStorage()
            filename = fs.save(uploaded_file.name, uploaded_file)
            file_path = os.path.join(fs.location, filename)

            # Extract text based on the file type
            if uploaded_file.name.endswith('.pdf'):
                extracted_text = extract_text_from_pdf(file_path)
            elif uploaded_file.name.endswith('.docx'):
                extracted_text = extract_text_from_docx(file_path)
            elif uploaded_file.name.endswith('.txt'):
                extracted_text = extract_text_from_txt(file_path)
            else:
                return JsonResponse({'message': 'Unsupported file type.', 'status': 'error'})

            # Append extracted text to the context list
            context_list.append(extracted_text)

        if chat_query is not None:
            bot_response = generate_prompt2(context_list,chat_query)
            return JsonResponse({'message': bot_response, 'status': 'success'})

    return render(request, 'BookChat.html')


# Initialize the Groq client
client = Groq(api_key="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")

# Define the maximum number of messages to retain in chat history
MAX_CHAT_HISTORY_LENGTH = 10


# Initialize the Groq client
client = Groq(api_key="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")

# Define the maximum number of messages to retain in chat history
MAX_CHAT_HISTORY_LENGTH = 10


import re



def format_response(response):
    """
    Format the assistant's response to match the MCQ layout with proper styling.
    Handles MCQs, regular questions, explanations, and HTML headings with consistent formatting.
    """
    # Clean the response and split into lines
    response = response.strip().replace('*', '')
    lines = response.split('\n')
    formatted_lines = []

    in_mcq = False  # Track if we're inside an MCQ block

    for line in lines:
        line = line.strip()  # Clean up each line

        # Skip empty lines to reduce unnecessary line breaks
        if not line:
            continue

        # Handle h1 headings
        if line.startswith('<h1>') and line.endswith('</h1>'):
            formatted_lines.append(f'<h1 class="text-center mb-4 text-2xl font-bold">{line[4:-5]}</h1>')

        # Handle h2 headings
        elif line.startswith('<h2>') and line.endswith('</h2>'):
            formatted_lines.append(f'<h2 class="text-center mb-3 text-xl font-semibold">{line[4:-5]}</h2>')

        # Handle h3 headings
        elif line.startswith('<h3>') and line.endswith('</h3>'):
            formatted_lines.append(f'<h3 class="text-center mb-2 text-lg font-medium">{line[4:-5]}</h3>')

        # Handle numbered questions (e.g., "1. What is...")
        elif re.match(r'^\d+\.', line):
            if in_mcq:
                formatted_lines.append('</div>')  # Close previous MCQ block if exists
            formatted_lines.append('<div class="question-block mb-3">')
            formatted_lines.append(f'<div class="question-text font-weight-bold">{line}</div>')
            in_mcq = True

        # Handle MCQ options (A), B), C), D))
        elif in_mcq and re.match(r'^[A-D]\)', line):
            formatted_lines.append(f'<div class="option ml-3">{line}</div>')

        # Handle blank answer spaces (e.g., "_____" or multiple underscores)
        elif '_' * 5 in line:
            formatted_lines.append('<div class="answer-space border-bottom my-2"></div>')

        # Handle regular text and explanations
        else:
            if in_mcq:
                formatted_lines.append('</div>')  # Close MCQ block
                in_mcq = False
            formatted_lines.append(f'<div class="content-text">{line}</div>')

    # Close any remaining open question div
    if in_mcq:
        formatted_lines.append('</div>')

    # Join all lines into a single formatted response with minimal line breaks
    formatted_response = ''.join(formatted_lines)
    return formatted_response

def chat(request):
    # Clear chat history on GET requests (page reload)
    if request.method == "GET":
        if "chat_history" in request.session:
            del request.session["chat_history"]
        return render(request, "chat.html")

    if request.method == "POST":
        user_input = request.POST.get("query", "")
        is_new_chat = request.POST.get("new_chat") == "true"

        # Initialize or clear chat history
        if is_new_chat or "chat_history" not in request.session:
            request.session["chat_history"] = []

        chat_history = request.session["chat_history"]
        chat_history.append({"role": "user", "content": user_input})

        if len(chat_history) > MAX_CHAT_HISTORY_LENGTH:
            chat_history = chat_history[-MAX_CHAT_HISTORY_LENGTH:]

        system_prompt = {
            "role": "system",
            "content": """You are a helpful assistant specializing in education and content generation. When responding to user queries:

                           1. For specific topics or concepts:
                              - Provide clear, concise explanations.
                              - Include relevant examples when possible.
                              - Highlight key points with bullet points.
                              - Do not add asteric in the it.

                           2. For multiple-choice questions (MCQs):
                              - Use numbered formatting.
                              - Follow this structure:
                              -DO not provide answers
                                1. Question text
                                   A) Option
                                   B) Option
                                   C) Option
                                   D) Option

                           3. For answers with explanations:
                              - Number each answer.
                              - Offer detailed explanations with examples where appropriate.
                              - Use clear paragraph breaks for readability.

                           4. For short questions requiring a written response:
                              - Number each question.
                              - Add answer spaces formatted like:
                                1. Question text
                                   _______________________
                                   _______________________

                                2. Question text
                                   _______________________
                                   _______________________

                              - Ensure adequate space for written answers with underscores.
                              - Format each question on a new line.
                           5. For differentiation questions:
                              - Use numbered formatting.
                              - Add answer spaces formatted like:
                               Topic 1 | topic 2
                                       |
                                 

                           Always follow consistent formatting and spacing to make responses easy to read."""
        }
        messages = [system_prompt] + chat_history

        try:
            chat_completion = client.chat.completions.create(
                messages=messages,
                model="llama3-8b-8192",
                temperature=0.5,
                max_tokens=1024,
            )

            response_content = chat_completion.choices[0].message.content
            formatted_response = format_response(response_content)

            chat_history.append({"role": "assistant", "content": response_content})
            request.session["chat_history"] = chat_history
            request.session.modified = True

            return JsonResponse({
                "response": formatted_response,
                "raw_response": response_content,
                "chat_history": chat_history
            })

        except Exception as e:
            print(f"Error in chat view: {e}")
            return JsonResponse({
                "error": f"An error occurred: {str(e)}",
                "chat_history": chat_history
            }, status=500)

    return JsonResponse({"error": "Method not allowed"}, status=405)
def paper(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            response_content = data.get("text", "")
            file_format = request.GET.get('format', '').lower()

            if not response_content:
                return JsonResponse({"error": "No content provided"}, status=400)

            if file_format == 'pdf':
                buffer = io.BytesIO()
                p = canvas.Canvas(buffer, pagesize=A4)
                width, height = A4
                margin = inch

                text_object = p.beginText(margin, height - margin)
                text_object.setFont("Helvetica", 11)
                text_object.setLeading(14)

                available_width = width - 2 * margin
                lines = response_content.split('\n')

                for line in lines:
                    if '_' * 5 in line:
                        current_y = text_object.getY()
                        p.line(margin, current_y - 5, width - margin, current_y - 5)
                        text_object.moveCursor(0, -20)
                    else:
                        words = line.split()
                        current_line = []
                        current_width = 0

                        for word in words:
                            word_width = p.stringWidth(word + ' ', "Helvetica", 11)
                            if current_width + word_width <= available_width:
                                current_line.append(word)
                                current_width += word_width
                            else:
                                text_object.textLine(' '.join(current_line))
                                current_line = [word]
                                current_width = word_width

                        if current_line:
                            text_object.textLine(' '.join(current_line))

                    if line.strip().endswith('?'):
                        text_object.moveCursor(0, -10)

                p.drawText(text_object)
                p.save()

                buffer.seek(0)
                response = HttpResponse(buffer, content_type='application/pdf')
                response['Content-Disposition'] = 'attachment; filename="response.pdf"'
                return response

            elif file_format == 'txt':
                buffer = io.BytesIO()
                buffer.write(response_content.encode('utf-8'))
                buffer.seek(0)
                response = HttpResponse(buffer, content_type='text/plain')
                response['Content-Disposition'] = 'attachment; filename="response.txt"'
                return response

            elif file_format == 'docx':
                buffer = io.BytesIO()
                doc = Document()
                style = doc.styles['Normal']
                style.font.name = 'Times New Roman'
                style.font.size = Pt(12)

                lines = response_content.split('\n')
                for line in lines:
                    if '_' * 5 in line:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_after = Pt(20)
                        p.paragraph_format.line_spacing = 2.0
                    else:
                        p = doc.add_paragraph(line)
                        if line.strip().endswith('?'):
                            p.paragraph_format.space_after = Pt(12)

                doc.save(buffer)
                buffer.seek(0)
                response = HttpResponse(
                    buffer,
                    content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )
                response['Content-Disposition'] = 'attachment; filename="response.docx"'
                return response

            else:
                return JsonResponse({"error": f"Unsupported format: {file_format}"}, status=400)

        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON data"}, status=400)
        except Exception as e:
            print(f"Error in paper view: {e}")
            return JsonResponse({"error": f"An error occurred: {str(e)}"}, status=500)

    return JsonResponse({"error": "Only POST method is allowed"}, status=405)

import requests


def ocr(request):
    if request.method == 'POST':
        # Decode the JSON data from the request
        data = json.loads(request.body)
        image_data = data.get('image_url')  # Expected to be a URL or Base64 data

        if image_data:
            # Check if `image_data` is a URL; if so, fetch and convert it to Base64
            if not image_data.startswith("data:image/"):
                response = requests.get(image_data)
                image_data = base64.b64encode(response.content).decode('utf-8')
                image_data = f"data:image/png;base64,{image_data}"  # Wrap in data URI format

            # Call the API with Base64 image data
            completion = client.chat.completions.create(
                model="llama-3.2-11b-vision-preview",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Display the text of the image.Do not give explanation."
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": image_data  # Send the Base64 data as a data URI
                                }
                            }
                        ]
                    }
                ],
                temperature=0.3,
                max_tokens=1024,
                top_p=1,
                stream=False,
                stop=None,
            )

            # Extract OCR result
            ocr_result = completion.choices[0].message.content

            # Return OCR result as JSON
            return JsonResponse({'extractedText': ocr_result})

    # For GET requests, render the upload page
    return render(request, 'ocr.html')
